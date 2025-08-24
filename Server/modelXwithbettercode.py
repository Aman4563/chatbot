from __future__ import annotations

import base64
import csv
import io
import json
import os
import re
from collections import Counter
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, Generator, List, Optional, Tuple

import dotenv
import docx
import PyPDF2
from huggingface_hub import InferenceClient

# NEW SDK (preferred)
from google.genai import types

from langchain_core.messages import AIMessage, HumanMessage, SystemMessage
from langchain_anthropic import ChatAnthropic
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_mistralai import ChatMistralAI
from langchain_openai import ChatOpenAI

from schemas import ChatRequest, FileData, Message as APIMessage

dotenv.load_dotenv()

# ==============================
# Cached environment keys
# ==============================
ENV_KEYS = {
    "OPENAI_API_KEY": os.getenv("OPENAI_API_KEY"),
    "GOOGLE_API_KEY": os.getenv("GOOGLE_API_KEY"),
    "MISTRAL_API_KEY": os.getenv("MISTRAL_API_KEY"),
    "ANTHROPIC_API_KEY": os.getenv("ANTHROPIC_API_KEY"),
    "HUGGINGFACE_API_KEY": os.getenv("HUGGINGFACE_API_KEY"),
}

# ==============================
# Model configuration (providers)
# ==============================
MODEL_CONFIG = {
    "mistral-large": {
        "provider": "mistral",
        "class": ChatMistralAI,
        "model": "mistral-large-latest",
        "supports_vision": False,
        "supports_files": True,
        "api_key": ENV_KEYS["MISTRAL_API_KEY"],
    },
    "mistral-small": {
        "provider": "mistral",
        "class": ChatMistralAI,
        "model": "mistral-small-latest",
        "supports_vision": False,
        "supports_files": True,
        "api_key": ENV_KEYS["MISTRAL_API_KEY"],
    },
    "gpt-4o": {
        "provider": "openai",
        "class": ChatOpenAI,
        "model": "gpt-4o",
        "supports_vision": True,
        "supports_files": True,
        "api_key": ENV_KEYS["OPENAI_API_KEY"],
    },
    "gpt-4o-mini": {
        "provider": "openai",
        "class": ChatOpenAI,
        "model": "gpt-4o-mini",
        "supports_vision": True,
        "supports_files": True,
        "api_key": ENV_KEYS["OPENAI_API_KEY"],
    },
    "claude-3-5-sonnet": {
        "provider": "anthropic",
        "class": ChatAnthropic,
        "model": "claude-3-5-sonnet-20241022",
        "supports_vision": True,
        "supports_files": True,
        "api_key": ENV_KEYS["ANTHROPIC_API_KEY"],
    },
    "gemini-1.5-pro": {
        "provider": "google",
        "class": ChatGoogleGenerativeAI,
        "model": "gemini-1.5-pro",
        "supports_vision": True,
        "supports_files": True,
        "api_key": ENV_KEYS["GOOGLE_API_KEY"],
    },
    "gemini-1.5-flash": {
        "provider": "google",
        "class": ChatGoogleGenerativeAI,
        "model": "gemini-1.5-flash",
        "supports_vision": True,
        "supports_files": True,
        "api_key": ENV_KEYS["GOOGLE_API_KEY"],
    },
    "gemini-2.0-flash": {
        "provider": "google",
        "class": ChatGoogleGenerativeAI,
        "model": "gemini-2.0-flash",
        "supports_vision": True,
        "supports_files": True,
        "api_key": ENV_KEYS["GOOGLE_API_KEY"],
    },
    "gemini-2.5-flash": {
        "provider": "google",
        "class": ChatGoogleGenerativeAI,  # FIXED (was a string)
        "model": "gemini-2.5-flash",
        "supports_vision": True,
        "supports_files": True,
        "api_key": ENV_KEYS["GOOGLE_API_KEY"],
    },
    "gemini-2.5-pro": {
        "provider": "google",
        "class": ChatGoogleGenerativeAI,
        "model": "gemini-2.5-pro",
        "supports_vision": True,
        "supports_files": True,
        "api_key": ENV_KEYS["GOOGLE_API_KEY"],
    },
}


# ================
# Optional web tool
# ================
def web_search_tool(query: str, num_results: int = 5) -> List[str]:
    """
    Simple web-search wrapper using googlesearch-python.
    Returns a list of absolute URLs (http/https).
    """
    try:
        from googlesearch import search
    except ImportError as e:
        raise RuntimeError("Missing dependency: pip install googlesearch-python") from e

    results: List[str] = []
    for url in search(query):
        if not url.lower().startswith(("http://", "https://")):
            continue
        results.append(url)
        if len(results) >= num_results:
            break
    return results


# ==========================================================
# REAL-TIME REQUIREMENT: LLM-JUDGED + HEURISTIC FALLBACK
# ==========================================================
# Lightweight heuristic primitives (used only as fallback)
_REL_TIME_PAT = re.compile(
    r"""
    \b(
        live|breaking|real[-\s]?time|updat(e|es|ed)|current|now|today|tonight|
        yesterday|tomorrow|this\s+(week|month|year|quarter)|
        last\s+\d+\s+(min|mins|minutes|hours|days|weeks|months|years)|
        past\s+\d+\s+(min|mins|minutes|hours|days|weeks|months|years)|
        recent|latest|trend(s|ing)?|status
    )\b
    """,
    re.IGNORECASE | re.VERBOSE,
)

_DATE_LIKE_PAT = re.compile(
    r"""
    (
        \b\d{1,2}[/.-]\d{1,2}([/.-]\d{2,4})?\b |      # 12/08[/2025], 12-8-25
        \b\d{4}-\d{1,2}-\d{1,2}\b |                  # 2025-08-18
        \b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)
           (?:uary|ch|e|y|e|e|y|ust|tember|ober|ember)?
           \s+\d{1,2}(?:,\s*\d{2,4})?\b
    )
    """,
    re.IGNORECASE | re.VERBOSE,
)

_DOMAIN_HINT_PAT = re.compile(
    r"""
    \b(
        price|prices|quote|quotes|stock|stocks|share|shares|ticker|market|index|
        exchange\s+rate|forex|currency|crypto|bitcoin|btc|eth|
        weather|forecast|temperature|
        score|scores|match|matches|fixture|fixtures|result|results|
        poll|polls|election|
        flight\s+status|train\s+status|air|rail|release\s+date|expected\s+to\s+release
    )\b
    """,
    re.IGNORECASE | re.VERBOSE,
)

_CURRENCY_SYMBOLS = {"$", "£", "€", "¥", "₹"}
_CURRENCY_CODES = {
    "USD", "EUR", "GBP", "JPY", "CNY", "INR", "AUD", "CAD", "CHF", "SEK",
    "NZD", "SGD", "HKD", "ZAR", "AED", "SAR"
}
_CRYPTO_CODES = {"BTC", "ETH", "SOL", "DOGE", "ADA", "XRP", "BNB", "USDT", "USDC"}

_COMPANY_SUFFIXES = {"ltd", "limited", "inc", "corp", "co", "plc", "llc", "ag", "oy", "sa", "nv", "ab"}
_COMMON_STOPWORDS = {
    "what", "is", "the", "a", "an", "to", "for", "of", "in", "and", "or", "on", "me", "give",
    "get", "show", "tell", "please", "pls", "now", "today", "how", "much", "latest", "current",
    "quote", "price", "rate", "when", "release", "date", "expected"
}
_TOKEN_PAT = re.compile(r"[A-Za-z][A-Za-z0-9\.'-]*|[%s]|[A-Z]{3}" % "".join(map(re.escape, _CURRENCY_SYMBOLS)))

# Precompute month names once
_MONTH_NAMES = tuple({
    datetime(2000, m, 1).strftime("%B").lower()
    for m in range(1, 13)
} | {
    datetime(2000, m, 1).strftime("%b").lower()
    for m in range(1, 13)
})


def _month_names() -> List[str]:
    return list(_MONTH_NAMES)


def _tokens(text: str) -> List[str]:
    return _TOKEN_PAT.findall(text or "")


def _has_currency(tokens: List[str]) -> bool:
    if any(t in _CURRENCY_SYMBOLS for t in tokens):
        return True
    if any(t.upper() in _CURRENCY_CODES for t in tokens):
        return True
    joined = " ".join(tokens).upper()
    if re.search(r"\b([A-Z]{3})/([A-Z]{3})\b", joined):
        return True
    if re.search(r"\b([A-Z]{6})\b", joined):  # EURUSD, USDINR
        return True
    return False


def _is_ticker_like(token: str) -> bool:
    if token.upper() in _CRYPTO_CODES:
        return True
    if re.fullmatch(r"[A-Z]{1,5}", token):
        return True
    return False


def _is_companyish(token: str) -> bool:
    t = token.lower().strip(".")
    return t in _COMPANY_SUFFIXES


def _candidate_entities(tokens: List[str]) -> List[str]:
    ents: List[str] = []
    n = len(tokens)
    for i, tok in enumerate(tokens):
        if tok.lower() in _COMMON_STOPWORDS:
            continue
        if tok in _CURRENCY_SYMBOLS or tok.upper() in _CURRENCY_CODES:
            continue
        if _is_ticker_like(tok):
            ents.append(tok)
            continue
        if re.fullmatch(r"[A-Za-z][A-Za-z0-9.'-]{1,14}", tok):
            ents.append(tok)
        if _is_companyish(tok) and i > 0:
            prev = tokens[i - 1]
            if re.fullmatch(r"[A-Za-z][A-Za-z0-9.'-]{1,20}", prev) and prev.lower() not in _COMMON_STOPWORDS:
                ents.append(prev)
    seen = set()
    out = []
    for e in ents:
        key = e.lower()
        if key not in seen:
            out.append(e)
            seen.add(key)
    return out


def _infer_history_topics(history: Optional[List[APIMessage]]) -> Counter:
    topics = Counter()
    if not history:
        return topics
    for msg in history:
        if getattr(msg, "content", None):
            t = msg.content.lower()
            if re.search(r"\b(stock|ticker|market|share|price|quote|nse|bse|nasdaq|nyse|forex|currency|crypto)\b", t):
                topics["finance"] += 1
            if re.search(r"\b(weather|forecast|temperature)\b", t):
                topics["weather"] += 1
            if re.search(r"\b(score|scores|match|fixture|result|team|league)\b", t):
                topics["sports"] += 1
            if re.search(r"\b(election|poll|votes?)\b", t):
                topics["politics"] += 1
            if re.search(r"\b(release|premiere|box\s*office|film|movie|episode|season)\b", t):
                topics["media_release"] += 1
    return topics


def _domain_from_signals(tokens: List[str], ents: List[str]) -> Optional[str]:
    joined = " ".join(tokens).lower()
    if _has_currency(tokens) or any(_is_ticker_like(t) for t in tokens) or any(_is_companyish(t) for t in tokens):
        return "finance"
    if re.search(r"\b(stock|ticker|market|share|price|quote|nse|bse|nasdaq|nyse|forex|currency|crypto)\b", joined):
        return "finance"
    if re.search(r"\b(weather|forecast|temperature)\b", joined):
        return "weather"
    if re.search(r"\b(score|scores|match|fixture|fixtures|result|results|odds|spread)\b", joined):
        return "sports"
    if re.search(r"\b(election|poll|votes?)\b", joined):
        return "politics"
    if re.search(r"\b(release\s+date|expected\s+to\s+release|premiere|film|movie|season|episode)\b", joined):
        return "media_release"
    if re.search(r"\b(flight|train)\s+status\b", joined):
        return "transport"
    return None


# -------- LLM INTENT DECIDER (primary) --------
def _select_intent_model():
    """
    Pick a lightweight, inexpensive model for intent classification.
    Preference order:
      - Gemini Flash (if GOOGLE_API_KEY)
      - GPT-4o-mini (if OPENAI_API_KEY)
      - Mistral-small (if MISTRAL_API_KEY)
      - Claude Sonnet (if ANTHROPIC_API_KEY)
    """
    if ENV_KEYS["GOOGLE_API_KEY"]:
        return ChatGoogleGenerativeAI(
            model="gemini-1.5-flash",
            google_api_key=ENV_KEYS["GOOGLE_API_KEY"],
            temperature=0,
            max_output_tokens=256,
        )
    if ENV_KEYS["OPENAI_API_KEY"]:
        return ChatOpenAI(
            model="gpt-4o-mini",
            api_key=ENV_KEYS["OPENAI_API_KEY"],
            temperature=0,
            max_tokens=256,
        )
    if ENV_KEYS["MISTRAL_API_KEY"]:
        return ChatMistralAI(
            model="mistral-small-latest",
            api_key=ENV_KEYS["MISTRAL_API_KEY"],
            temperature=0,
            max_tokens=256,
        )
    if ENV_KEYS["ANTHROPIC_API_KEY"]:
        return ChatAnthropic(
            model="claude-3-5-sonnet-20241022",
            api_key=ENV_KEYS["ANTHROPIC_API_KEY"],
            temperature=0,
            max_tokens=256,
        )
    return None  # no keys


_INTENT_SYSTEM = """You are an Intent Decider.
Decide if answering the user's request correctly requires up-to-date information from the public web.
Be conservative with false negatives: if it is plausible the correct answer depends on recent events, releases, prices, schedules, or status, return true.
Output STRICT JSON:
{"needs_realtime": true|false, "category": "<one of: finance, weather, sports, politics, transport, media_release, outage_status, general>", "confidence": 0..100}
No extra text.
"""

_INTENT_USER_TEMPLATE = """User request:
{prompt}

Short conversation context (optional, may be empty):
{history}

Decide whether web/real-time info is needed. Examples:
- "price for astral" => needs_realtime true, category "finance"
- "when is Avengers Doomsday expected to release?" => true, "media_release"
- "weather in Delhi tomorrow" => true, "weather"
- "what is photosynthesis?" => false, "general"
Return JSON only."""


def _history_compact(history: Optional[List[APIMessage]]) -> str:
    if not history:
        return ""
    keep = []
    for m in history[-6:]:
        role = m.role
        text = (m.content or "").strip().replace("\n", " ")
        if len(text) > 200:
            text = text[:200] + "..."
        keep.append(f"{role}: {text}")
    return "\n".join(keep)


def _llm_decide_realtime(prompt: str, history: Optional[List[APIMessage]]) -> Optional[Dict[str, Any]]:
    model = _select_intent_model()
    if not model:
        return None
    try:
        sys = SystemMessage(content=_INTENT_SYSTEM)
        user = HumanMessage(content=_INTENT_USER_TEMPLATE.format(prompt=prompt or "", history=_history_compact(history)))
        resp = model.invoke([sys, user])
        raw = getattr(resp, "content", "") or ""
        # Some models return a list of content parts; normalize:
        if isinstance(raw, list):
            raw = " ".join([str(x) for x in raw])
        data = json.loads(raw)
        if not isinstance(data, dict):
            return None
        # Validate fields
        needs = bool(data.get("needs_realtime", False))
        cat = data.get("category", "general") or "general"
        conf = int(data.get("confidence", 0)) if isinstance(data.get("confidence", 0), (int, float)) else 0
        return {"needs_realtime": needs, "category": cat, "confidence": conf}
    except Exception:
        return None


# -------- Heuristic fallback (secondary) --------
def analyze_realtime_intent_fallback(prompt: str, history: Optional[List[APIMessage]] = None) -> Dict[str, Any]:
    text = (prompt or "").strip()
    tokens = _tokens(text)
    ents = _candidate_entities(tokens)

    score = 0
    signals: Dict[str, Any] = {}
    domain = _domain_from_signals(tokens, ents)

    if domain:
        score += 2
    if _has_currency(tokens):
        score += 2
        signals["currency"] = True
    if any(_is_ticker_like(t) for t in tokens):
        score += 2
        signals["ticker_like"] = True
    if ents:
        score += 1
        signals["entities"] = ents

    token_count = sum(bool(re.search(r"[A-Za-z0-9]", t)) for t in tokens)
    if token_count <= 3 and (signals.get("ticker_like") or domain in {"finance", "media_release"} or ents):
        score += 1
        signals["short_query_bias"] = True

    if _REL_TIME_PAT.search(text):
        score += 1
        signals["rel_time"] = True
    if _DATE_LIKE_PAT.search(text):
        score += 1
        signals["date_like"] = True
    if _DOMAIN_HINT_PAT.search(text):
        score += 1
        signals["domain_hint_words"] = True

    year = datetime.now().year
    if str(year) in text:
        score += 1
        signals["mentions_year"] = True
    if any(m in text.lower() for m in _MONTH_NAMES):
        score += 1
        signals["mentions_month"] = True

    hist = _infer_history_topics(history)
    if hist:
        top_topic, cnt = hist.most_common(1)[0]
        signals["history_topic"] = {top_topic: cnt}
        if domain and top_topic == domain:
            score += 2
        elif not domain and cnt >= 2:
            domain = top_topic
            score += 1

    needs_realtime = score >= 3
    return {
        "needs_realtime": needs_realtime,
        "category": domain or "general",
        "score": score,
        "signals": signals,
    }


def analyze_realtime_intent(prompt: str, history: Optional[List[APIMessage]] = None) -> Dict[str, Any]:
    # Try semantic LLM judge first
    judged = _llm_decide_realtime(prompt, history)
    if judged is not None:
        return {
            "needs_realtime": bool(judged.get("needs_realtime", False)),
            "category": judged.get("category", "general"),
            "score": judged.get("confidence", 0),
            "signals": {"judge": "llm", "confidence": judged.get("confidence", 0)},
        }
    # Fallback if LLM judge unavailable/fails
    return analyze_realtime_intent_fallback(prompt, history)


def needs_realtime_data(prompt: str, history: Optional[List[APIMessage]] = None) -> bool:
    return analyze_realtime_intent(prompt, history).get("needs_realtime", False)


# Phrases that explicitly request web search (forces tool on)
_FORCE_SEARCH_TERMS = (
    "search", "search online", "google it", "google this", "lookup", "look up",
    "browse", "check online", "check the web", "web search"
)


def _wants_forced_search(prompt: str) -> bool:
    p = (prompt or "").strip().lower()
    return any(term in p for term in _FORCE_SEARCH_TERMS)


def get_gemini_tools_if_needed(model_name: str, prompt: str, history: Optional[List[APIMessage]] = None):
    """
    If using a Gemini model and the prompt likely needs real-time info,
    attach Google Search grounding tool (new google.genai SDK).
    Returns a list (possibly empty) compatible with LangChain's ChatGoogleGenerativeAI.
    """
    if not (model_name or "").startswith("gemini"):
        return []

    if not (_wants_forced_search(prompt) or needs_realtime_data(prompt, history)):
        return []

    # Preferred: new SDK Tool object (passes through in newer LangChain versions)
    try:
        return [types.Tool(google_search=types.GoogleSearch())]
    except Exception:
        # Compatibility fallback for wrappers that expect dicts
        return [{"google_search": {}}]


# =================
# Image generation
# =================
def image_gen_tool(prompt: str) -> str:
    """
    Image generation via Hugging Face Inference API.
    Returns a data URL (base64 PNG) or a stable placeholder URL if generation fails.
    """
    hf_token = ENV_KEYS["HUGGINGFACE_API_KEY"]
    if hf_token:
        try:
            client = InferenceClient(api_key=hf_token)
            model_id = os.getenv("HF_TTI_MODEL", "black-forest-labs/FLUX.1-schnell")
            image = client.text_to_image(prompt, model=model_id)

            buf = BytesIO()
            image.save(buf, format="PNG")
            img_bytes = buf.getvalue()
            b64 = base64.b64encode(img_bytes).decode("utf-8")
            return f"data:image/png;base64,{b64}"
        except Exception as e:
            print(f"Hugging Face image generation failed: {e!r}")
    return "https://picsum.photos/512"


# =============================
# Model initialization helpers
# =============================
def get_model(model_name: str):
    """Initialize and return the appropriate LangChain chat model."""
    if model_name not in MODEL_CONFIG:
        raise ValueError(f"Unsupported model: {model_name}")

    cfg = MODEL_CONFIG[model_name]
    api_key = cfg["api_key"]
    if not api_key:
        raise ValueError(f"API key not found for {cfg['provider']} models")

    provider = cfg["provider"]
    model_cls = cfg["class"]
    common = dict(temperature=0.7, streaming=True)

    if provider == "google":
        # ChatGoogleGenerativeAI expects google_api_key and max_output_tokens
        return model_cls(
            model=cfg["model"],
            google_api_key=api_key,
            max_output_tokens=8192,
            **common,
        )
    else:
        # OpenAI, Mistral, Anthropic: use api_key + max_tokens
        return model_cls(
            model=cfg["model"],
            api_key=api_key,
            max_tokens=8192,
            **common,
        )


# ==========================
# File & content processing
# ==========================
def process_image_for_langchain(file_data: FileData) -> Optional[Dict[str, Any]]:
    """Prepare an image attachment for multimodal models that accept image_url."""
    try:
        return {
            "type": "image_url",
            "image_url": {"url": f"data:{file_data.mime_type};base64,{file_data.data}"},
        }
    except Exception as e:
        print(f"Error processing image {file_data.filename}: {e}")
        return None


def extract_pdf_content(file_data: FileData) -> str:
    """Extract text content from a PDF."""
    try:
        decoded = base64.b64decode(file_data.data)
        pdf_file = io.BytesIO(decoded)
        reader = PyPDF2.PdfReader(pdf_file)
        out = []
        for i, page in enumerate(reader.pages, 1):
            try:
                text = page.extract_text() or ""
            except Exception as e:
                text = f"[Error extracting page {i}: {e}]"
            if text.strip():
                out.append(f"--- Page {i} ---\n{text}")
        return "\n".join(out) if out else "PDF appears empty or image-only."
    except Exception as e:
        return f"Error extracting PDF: {e}"


def extract_docx_content(file_data: FileData) -> str:
    """Extract text content from a DOCX file."""
    try:
        decoded = base64.b64decode(file_data.data)
        docx_file = io.BytesIO(decoded)
        document = docx.Document(docx_file)
        parts: List[str] = []
        for p in document.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in document.tables:
            parts.append("\n--- Table ---")
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts) if parts else "Word document appears to be empty."
    except Exception as e:
        return f"Error extracting Word document content: {e}"


def extract_csv_content(file_data: FileData) -> str:
    """Extract and lightly format CSV content (first 100 rows)."""
    try:
        decoded = base64.b64decode(file_data.data).decode("utf-8")
        csv_reader = csv.reader(io.StringIO(decoded))
        lines: List[str] = ["CSV Data:"]
        for idx, row in enumerate(csv_reader, 1):
            if idx == 1:
                lines.append("Headers: " + " | ".join(row))
                lines.append("-" * 50)
            else:
                lines.append(f"Row {idx - 1}: " + " | ".join(row))
            if idx >= 100:
                lines.append("... (showing first 100 rows)")
                break
        return "\n".join(lines)
    except Exception as e:
        return f"Error processing CSV: {e}"


def extract_json_content(file_data: FileData) -> str:
    """Pretty-print JSON content."""
    try:
        decoded = base64.b64decode(file_data.data).decode("utf-8")
        obj = json.loads(decoded)
        return "JSON Data:\n" + json.dumps(obj, indent=2, ensure_ascii=False)
    except json.JSONDecodeError as e:
        return f"Invalid JSON: {e}"
    except Exception as e:
        return f"Error processing JSON: {e}"


def process_document_for_langchain(file_data: FileData) -> str:
    """Route to the right extractor for supported document types."""
    try:
        if file_data.mime_type == "text/plain":
            decoded = base64.b64decode(file_data.data).decode("utf-8")
            return f"Text Document: {file_data.filename}\n\n{decoded}"
        if file_data.mime_type == "text/csv":
            return f"CSV Document: {file_data.filename}\n\n{extract_csv_content(file_data)}"
        if file_data.mime_type == "application/json":
            return f"JSON Document: {file_data.filename}\n\n{extract_json_content(file_data)}"
        if file_data.mime_type == "application/pdf":
            return f"PDF Document: {file_data.filename}\n\nExtracted Content:\n{extract_pdf_content(file_data)}"
        if file_data.mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return f"Word Document: {file_data.filename}\n\nExtracted Content:\n{extract_docx_content(file_data)}"
        if file_data.mime_type == "application/msword":
            return (
                f"Legacy Word Document: {file_data.filename}\n\n"
                "[Legacy .doc requires conversion to .docx for full analysis]"
            )
        return f"Document: {file_data.filename} (Type: {file_data.mime_type})\n[Unsupported type]"
    except Exception as e:
        return f"Error processing document {file_data.filename}: {e}"


# ======================
# Conversation assembly
# ======================
def build_conversation_messages(
    history: List[APIMessage], current: ChatRequest, system_prompt: Optional[str]
) -> List:
    """Compose messages for LangChain models, including multimodal content."""
    msgs: List[Any] = []
    if system_prompt:
        msgs.append(SystemMessage(content=system_prompt))

    # Past turns
    for msg in history:
        if msg.role == "user":
            content_parts: List[Dict[str, Any]] = []
            if msg.content:
                content_parts.append({"type": "text", "text": msg.content})
            if getattr(msg, "files", None):
                for f in msg.files:
                    if f.mime_type.startswith("image/"):
                        img = process_image_for_langchain(f)
                        if img:
                            content_parts.append(img)
                    else:
                        content_parts.append({"type": "text", "text": process_document_for_langchain(f)})
            if content_parts:
                if len(content_parts) == 1 and content_parts[0].get("type") == "text":
                    msgs.append(HumanMessage(content=content_parts[0]["text"]))
                else:
                    msgs.append(HumanMessage(content=content_parts))
        elif msg.role == "assistant":
            msgs.append(AIMessage(content=msg.content))

    # Current turn
    parts: List[Dict[str, Any]] = []
    if current.message.text:
        parts.append({"type": "text", "text": current.message.text})
    for f in current.message.files:
        if f.mime_type.startswith("image/"):
            img = process_image_for_langchain(f)
            if img:
                parts.append(img)
        else:
            parts.append({"type": "text", "text": process_document_for_langchain(f)})
    if parts:
        if len(parts) == 1 and parts[0].get("type") == "text":
            msgs.append(HumanMessage(content=parts[0]["text"]))
        else:
            msgs.append(HumanMessage(content=parts))
    return msgs


# ===================
# Inference routines
# ===================
def _invoke_with_optional_tools(model, messages, tools):
    # Try both 'tools' and 'tool_config' to support wrapper variations
    try:
        if tools:
            return model.invoke(messages, tools=tools)
        return model.invoke(messages)
    except TypeError:
        if tools:
            try:
                return model.invoke(messages, tool_config={"tools": tools})
            except TypeError:
                pass
        return model.invoke(messages)


def _stream_with_optional_tools(model, messages, tools):
    # Try both 'tools' and 'tool_config' to support wrapper variations
    try:
        if tools:
            return model.stream(messages, tools=tools)
        return model.stream(messages)
    except TypeError:
        if tools:
            try:
                return model.stream(messages, tool_config={"tools": tools})
            except TypeError:
                pass
        return model.stream(messages)


def my_genai_chat_function_stream(payload: ChatRequest) -> Generator[str, None, None]:
    """Streaming chat generation with optional Gemini web-search tool injection."""
    if not payload.message.text and not payload.message.files:
        yield "No message content provided."
        return
    try:
        model = get_model(payload.model_name)
        messages = build_conversation_messages(payload.history, payload, payload.system_prompt)
        if not messages:
            yield "No valid content to process."
            return
        tools = get_gemini_tools_if_needed(payload.model_name, payload.message.text or "", payload.history)
        for chunk in _stream_with_optional_tools(model, messages, tools):
            if getattr(chunk, "content", None):
                # chunk.content may be a string or a list of parts; normalize to string
                if isinstance(chunk.content, list):
                    yield " ".join(str(p) for p in chunk.content if p)
                else:
                    yield chunk.content
    except Exception as e:
        print(f"Error in LangChain stream call: {e}")
        yield f"Error generating response: {e}"


def my_genai_chat_function(payload: ChatRequest) -> str:
    """Non-streaming variant of chat generation."""
    if not payload.message.text and not payload.message.files:
        return "No message content provided."
    try:
        model = get_model(payload.model_name)
        messages = build_conversation_messages(payload.history, payload, payload.system_prompt)
        if not messages:
            return "No valid content to process."
        tools = get_gemini_tools_if_needed(payload.model_name, payload.message.text or "", payload.history)
        resp = _invoke_with_optional_tools(model, messages, tools)
        content = getattr(resp, "content", None)
        if isinstance(content, list):
            return " ".join(str(p) for p in content if p) or "No response from model."
        return content or "No response from model."
    except Exception as e:
        print(f"Error in LangChain call: {e}")
        return f"Error generating response: {e}"
