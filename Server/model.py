from __future__ import annotations

import base64
import csv
import io
import json
import os
import re
from io import BytesIO
from typing import Any, Dict, Generator, List, Optional, Tuple

import dotenv
import docx
import PyPDF2
from huggingface_hub import InferenceClient

# New Google SDK import kept (not used directly for tools due to wrapper limits)
from google.genai import types  # noqa: F401

from langchain_core.messages import AIMessage, HumanMessage, SystemMessage
from langchain_anthropic import ChatAnthropic
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_mistralai import ChatMistralAI
from langchain_openai import ChatOpenAI

from schemas import ChatRequest, FileData, Message as APIMessage

dotenv.load_dotenv()

# ==============================
# Cached environment keys / toggles
# ==============================
ENV_KEYS = {
    "OPENAI_API_KEY": os.getenv("OPENAI_API_KEY"),
    "GOOGLE_API_KEY": os.getenv("GOOGLE_API_KEY"),
    "MISTRAL_API_KEY": os.getenv("MISTRAL_API_KEY"),
    "ANTHROPIC_API_KEY": os.getenv("ANTHROPIC_API_KEY"),
    "HUGGINGFACE_API_KEY": os.getenv("HUGGINGFACE_API_KEY"),
}

def _get_env_flag(name: str, default: str) -> str:
    v = os.getenv(name)
    return v.strip().lower() if isinstance(v, str) and v.strip() else default

def _get_env_int(name: str, default: int) -> int:
    try:
        return int(os.getenv(name, "").strip())
    except Exception:
        return default

# Explicit control knobs
GEMINI_TOOL_DEFAULT = _get_env_flag("GEMINI_TOOL_DEFAULT", "auto")  # auto|always|never
GEMINI_TOOL_MAX_QUERIES = _get_env_int("GEMINI_TOOL_MAX_QUERIES", 2)
GEMINI_TOOL_DISABLE_FOR_CS = _get_env_flag("GEMINI_TOOL_DISABLE_FOR_CS", "1") in {"1", "true", "yes"}
GEMINI_IMAGE_TOOL_DEFAULT = _get_env_flag("GEMINI_IMAGE_TOOL_DEFAULT", "auto")  # auto|always|never

# ==============================
# Model configuration (providers)
# ==============================
MODEL_CONFIG: Dict[str, Dict[str, Any]] = {
    "mistral-large": {
        "provider": "mistral",
        "class": ChatMistralAI,
        "model": "mistral-large-latest",
        "supports_vision": False,
        "supports_files": True,
        "api_key": ENV_KEYS["MISTRAL_API_KEY"],
    },
    "mistral-small": {
        "provider": "mistral",
        "class": ChatMistralAI,
        "model": "mistral-small-latest",
        "supports_vision": False,
        "supports_files": True,
        "api_key": ENV_KEYS["MISTRAL_API_KEY"],
    },
    "gpt-4o": {
        "provider": "openai",
        "class": ChatOpenAI,
        "model": "gpt-4o",
        "supports_vision": True,
        "supports_files": True,
        "api_key": ENV_KEYS["OPENAI_API_KEY"],
    },
    "gpt-4o-mini": {
        "provider": "openai",
        "class": ChatOpenAI,
        "model": "gpt-4o-mini",
        "supports_vision": True,
        "supports_files": True,
        "api_key": ENV_KEYS["OPENAI_API_KEY"],
    },
    "claude-3-5-sonnet": {
        "provider": "anthropic",
        "class": ChatAnthropic,
        "model": "claude-3-5-sonnet-20241022",
        "supports_vision": True,
        "supports_files": True,
        "api_key": ENV_KEYS["ANTHROPIC_API_KEY"],
    },
    "gemini-1.5-pro": {
        "provider": "google",
        "class": ChatGoogleGenerativeAI,
        "model": "gemini-1.5-pro",
        "supports_vision": True,
        "supports_files": True,
        "api_key": ENV_KEYS["GOOGLE_API_KEY"],
    },
    "gemini-1.5-flash": {
        "provider": "google",
        "class": ChatGoogleGenerativeAI,
        "model": "gemini-1.5-flash",
        "supports_vision": True,
        "supports_files": True,
        "api_key": ENV_KEYS["GOOGLE_API_KEY"],
    },
    "gemini-2.0-flash": {
        "provider": "google",
        "class": ChatGoogleGenerativeAI,
        "model": "gemini-2.0-flash",
        "supports_vision": True,
        "supports_files": True,
        "api_key": ENV_KEYS["GOOGLE_API_KEY"],
    },
    "gemini-2.5-flash": {
        "provider": "google",
        "class": ChatGoogleGenerativeAI,
        "model": "gemini-2.5-flash",
        "supports_vision": True,
        "supports_files": True,
        "api_key": ENV_KEYS["GOOGLE_API_KEY"],
    },
    "gemini-2.5-pro": {
        "provider": "google",
        "class": ChatGoogleGenerativeAI,
        "model": "gemini-2.5-pro",
        "supports_vision": True,
        "supports_files": True,
        "api_key": ENV_KEYS["GOOGLE_API_KEY"],
    },
}

# ================
# Optional legacy web tool (not used by Gemini; kept for completeness)
# ================
def web_search_tool(query: str, num_results: int = 5) -> List[str]:
    try:
        from googlesearch import search
    except ImportError as e:
        raise RuntimeError("Missing dependency: pip install googlesearch-python") from e
    results: List[str] = []
    for url in search(query):
        if not url.lower().startswith(("http://", "https://")):
            continue
        results.append(url)
        if len(results) >= num_results:
            break
    return results

# ==========================
# File & content processing
# ==========================
def process_image_for_langchain(file_data: FileData) -> Optional[Dict[str, Any]]:
    """Prepare an image attachment for multimodal models that accept image_url."""
    try:
        return {
            "type": "image_url",
            "image_url": {"url": f"data:{file_data.mime_type};base64,{file_data.data}"},
        }
    except Exception as e:
        print(f"Error processing image {file_data.filename}: {e}")
        return None

def extract_pdf_content(file_data: FileData) -> str:
    """Extract text content from a PDF."""
    try:
        decoded = base64.b64decode(file_data.data)
        pdf_file = io.BytesIO(decoded)
        reader = PyPDF2.PdfReader(pdf_file)
        out = []
        for i, page in enumerate(reader.pages, 1):
            try:
                text = page.extract_text() or ""
            except Exception as e:
                text = f"[Error extracting page {i}: {e}]"
            if text.strip():
                out.append(f"--- Page {i} ---\n{text}")
        return "\n".join(out) if out else "PDF appears empty or image-only."
    except Exception as e:
        return f"Error extracting PDF: {e}"

def extract_docx_content(file_data: FileData) -> str:
    """Extract text content from a DOCX file."""
    try:
        decoded = base64.b64decode(file_data.data)
        docx_file = io.BytesIO(decoded)
        document = docx.Document(docx_file)
        parts: List[str] = []
        for p in document.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in document.tables:
            parts.append("\n--- Table ---")
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts) if parts else "Word document appears to be empty."
    except Exception as e:
        return f"Error extracting Word document content: {e}"

def extract_csv_content(file_data: FileData) -> str:
    """Extract and lightly format CSV content (first 100 rows)."""
    try:
        decoded = base64.b64decode(file_data.data).decode("utf-8")
        csv_reader = csv.reader(io.StringIO(decoded))
        lines: List[str] = ["CSV Data:"]
        for idx, row in enumerate(csv_reader, 1):
            if idx == 1:
                lines.append("Headers: " + " | ".join(row))
                lines.append("-" * 50)
            else:
                lines.append(f"Row {idx - 1}: " + " | ".join(row))
            if idx >= 100:
                lines.append("... (showing first 100 rows)")
                break
        return "\n".join(lines)
    except Exception as e:
        return f"Error processing CSV: {e}"

def extract_json_content(file_data: FileData) -> str:
    """Pretty-print JSON content."""
    try:
        decoded = base64.b64decode(file_data.data).decode("utf-8")
        obj = json.loads(decoded)
        return "JSON Data:\n" + json.dumps(obj, indent=2, ensure_ascii=False)
    except json.JSONDecodeError as e:
        return f"Invalid JSON: {e}"
    except Exception as e:
        return f"Error processing JSON: {e}"

def process_document_for_langchain(file_data: FileData) -> str:
    """Route to the right extractor for supported document types."""
    try:
        if file_data.mime_type == "text/plain":
            decoded = base64.b64decode(file_data.data).decode("utf-8")
            return f"Text Document: {file_data.filename}\n\n{decoded}"
        if file_data.mime_type == "text/csv":
            return f"CSV Document: {file_data.filename}\n\n{extract_csv_content(file_data)}"
        if file_data.mime_type == "application/json":
            return f"JSON Document: {file_data.filename}\n\n{extract_json_content(file_data)}"
        if file_data.mime_type == "application/pdf":
            return f"PDF Document: {file_data.filename}\n\nExtracted Content:\n{extract_pdf_content(file_data)}"
        if file_data.mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return f"Word Document: {file_data.filename}\n\nExtracted Content:\n{extract_docx_content(file_data)}"
        if file_data.mime_type == "application/msword":
            return (
                f"Legacy Word Document: {file_data.filename}\n\n"
                "[Legacy .doc requires conversion to .docx for full analysis]"
            )
        return f"Document: {file_data.filename} (Type: {file_data.mime_type})\n[Unsupported type]"
    except Exception as e:
        return f"Error processing document {file_data.filename}: {e}"

# ======================
# Conversation assembly
# ======================
def build_conversation_messages(
    history: List[APIMessage], current: ChatRequest, system_prompt: Optional[str]
) -> List:
    """Compose messages for LangChain models, including multimodal content."""
    msgs: List[Any] = []
    if system_prompt:
        msgs.append(SystemMessage(content=system_prompt))

    # Past turns
    for msg in history:
        if msg.role == "user":
            content_parts: List[Dict[str, Any]] = []
            if msg.content:
                content_parts.append({"type": "text", "text": msg.content})
            if getattr(msg, "files", None):
                for f in msg.files:
                    if f.mime_type.startswith("image/"):
                        img = process_image_for_langchain(f)
                        if img:
                            content_parts.append(img)
                    else:
                        content_parts.append({"type": "text", "text": process_document_for_langchain(f)})
            if content_parts:
                if len(content_parts) == 1 and content_parts[0].get("type") == "text":
                    msgs.append(HumanMessage(content=content_parts[0]["text"]))
                else:
                    msgs.append(HumanMessage(content=content_parts))
        elif msg.role == "assistant":
            msgs.append(AIMessage(content=msg.content))

    # Current turn
    parts: List[Dict[str, Any]] = []
    if current.message.text:
        parts.append({"type": "text", "text": current.message.text})
    for f in current.message.files:
        if f.mime_type.startswith("image/"):
            img = process_image_for_langchain(f)
            if img:
                parts.append(img)
        else:
            parts.append({"type": "text", "text": process_document_for_langchain(f)})
    if parts:
        if len(parts) == 1 and parts[0].get("type") == "text":
            msgs.append(HumanMessage(content=parts[0]["text"]))
        else:
            msgs.append(HumanMessage(content=parts))
    return msgs

# ===========================
# Gemini tool policy & control
# ===========================
_FORCE_SEARCH_TERMS = (
    "search online", "google it", "google this", "check the web",
    "web search", "browse the web", "look it up", "lookup online", "check online"
)

_CS_SEARCH_NEG = (
    "binary search", "search tree", "graph search", "search algorithm",
    "search complexity", "dfs", "bfs"
)

_FORCE_IMAGE_TERMS = (
    "generate an image", "generate image", "create an image", "make an image",
    "draw", "illustration", "diagram", "sketch", "logo", "wallpaper",
    "poster", "render", "picture of", "image of", "make a picture"
)

def _should_block_for_cs(prompt: str) -> bool:
    p = (prompt or "").lower()
    return GEMINI_TOOL_DISABLE_FOR_CS and any(kw in p for kw in _CS_SEARCH_NEG)

def _is_forced_search(prompt: str) -> bool:
    p = (prompt or "").lower()
    return any(kw in p for kw in _FORCE_SEARCH_TERMS)

def _is_forced_image(prompt: str) -> bool:
    p = (prompt or "").lower()
    return any(kw in p for kw in _FORCE_IMAGE_TERMS)

def _is_gemini_model(model_name: str) -> bool:
    return (model_name or "").startswith("gemini")

def _gemini_policy_text(max_queries: int, enable_image_json: bool) -> str:
    base = (
        "You can use web search (Google Search tool) ONLY if the answer depends on current or external information "
        "such as recent news, live results, release dates, prices, schedules, weather, or fast-changing facts. "
        f"Be frugal: perform at most {max_queries} searches unless absolutely necessary. "
        "If the query is answerable from your internal knowledge, do not use the tool. "
        "When you do use the tool, prefer a small number of targeted searches."
    )
    if enable_image_json:
        base += (
            "\nIf creating an image would satisfy the user's request or significantly help, "
            "emit a single-line JSON directive as the VERY FIRST LINE of your reply, like:\n"
            '{"tool":"generate_image","prompt":"<clear descriptive prompt>","model":"<optional hf model id>"}\n'
            "After that JSON line, continue with your normal answer."
        )
    return base

def get_gemini_tools_and_policy(model_name: str, prompt: str) -> Tuple[List[Any], Optional[str], bool]:
    """
    Build the Gemini tools list (only google_search to avoid function_declarations) and a policy message.
    Returns (tools, policy_text, image_json_enabled).
    image_json_enabled=True means we asked the model to optionally emit the image JSON directive.
    """
    if not _is_gemini_model(model_name):
        return [], None, False

    tools: List[Any] = []
    image_json_enabled = False

    # GOOGLE SEARCH TOOL (dict shape for wrapper compatibility)
    if GEMINI_TOOL_DEFAULT != "never" and not _should_block_for_cs(prompt):
        if _is_forced_search(prompt) or GEMINI_TOOL_DEFAULT in {"auto", "always"}:
            tools.append({"google_search": {}})

    # IMAGE directive (no function_declarations — we parse JSON ourselves)
    if GEMINI_IMAGE_TOOL_DEFAULT != "never" or _is_forced_image(prompt):
        image_json_enabled = True

    policy = _gemini_policy_text(GEMINI_TOOL_MAX_QUERIES, enable_image_json=image_json_enabled)
    return tools, policy, image_json_enabled

def _insert_policy_message(messages: List, policy: Optional[str]) -> List:
    if not policy:
        return messages
    idx = 0
    if messages and isinstance(messages[0], SystemMessage):
        idx = 1
    return messages[:idx] + [SystemMessage(content=policy)] + messages[idx:]

# =============================
# Model initialization helpers
# =============================
def get_model(model_name: str):
    """Initialize and return the appropriate LangChain chat model."""
    if model_name not in MODEL_CONFIG:
        raise ValueError(f"Unsupported model: {model_name}")

    cfg = MODEL_CONFIG[model_name]
    api_key = cfg["api_key"]
    if not api_key:
        raise ValueError(f"API key not found for {cfg['provider']} models")

    provider = cfg["provider"]
    model_cls = cfg["class"]

    # Do NOT pass streaming=True to ChatGoogleGenerativeAI (your build rejects it).
    if provider == "google":
        return model_cls(
            model=cfg["model"],
            google_api_key=api_key,
            max_output_tokens=8192,
            temperature=0.7,
        )
    else:
        return model_cls(
            model=cfg["model"],
            api_key=api_key,
            max_tokens=8192,
            temperature=0.7,
        )

# ===================
# Image generation
# ===================
# def image_gen_tool(prompt: str, model_id: Optional[str] = None) -> str:
#     """
#     Image generation via Hugging Face Inference API.
#     Returns a data URL (base64 PNG) or a stable placeholder URL if generation fails.
#     """
#     hf_token = ENV_KEYS["HUGGINGFACE_API_KEY"]
#     if hf_token:
#         try:
#             client = InferenceClient(api_key=hf_token)
#             image = client.text_to_image(
#                 prompt,
#                 model="ZB-Tech/Text-to-Image",
#             )
#             # print(f"Generated image: {image}")
#             # Convert PIL.Image to PNG bytes
#             buf = BytesIO()
#             image.save(buf, format="PNG")
#             buf.seek(0)
#             img_bytes = buf.getvalue()
#             # Base64-encode and return a data URL
#             b64 = base64.b64encode(img_bytes).decode("utf-8")
#             print(f"Generated image URL: data:image/png;base64,{b64}")
#             open("generated_image.png", "wb").write(img_bytes)  # Save for debugging
#             return f"data:image/png;base64,{b64}"
#         except Exception as e:
#             print(f"Hugging Face image generation failed: {e!r}")
#     return "https://picsum.photos/1024"


def image_gen_tool(prompt: str, model_id: Optional[str] = None, local_image_path: Optional[str] = None) -> str:
    """
    Image URL provider for frontend testing.

    Modes (set with env IMAGE_GEN_MODE):
      - local (default): read a local PNG and return data URL (no HF hits).
          * path comes from param `local_image_path` or env IMAGE_GEN_LOCAL_PATH
            (defaults to "generated_image.png")
      - hf: (kept commented out) Hugging Face generation (enable later if needed)
      - placeholder: return a static placeholder URL

    Returns:
      str: data:image/png;base64,...  (or an https placeholder if not available)
    """
    import os
    import base64
    from io import BytesIO

    mode = os.getenv("IMAGE_GEN_MODE", "local").strip().lower()  # local | hf | placeholder

    if mode == "local":
        # Prefer explicit arg, then env, then default
        path = local_image_path or os.getenv("IMAGE_GEN_LOCAL_PATH", "generated_image.png")
        try:
            with open(path, "rb") as f:
                img_bytes = f.read()
            # Optional: write a copy for debug (just to confirm we could read)
            with open("debug_local_image_copy.png", "wb") as dbg:
                dbg.write(img_bytes)
            b64 = base64.b64encode(img_bytes).decode("utf-8")
            print(f"Generated image URL: data:image/png;base64,{b64}")
            return f"data:image/png;base64,{b64}"
        except FileNotFoundError:
            print(f"[image_gen_tool] Local image not found at: {path!r}. Falling back to placeholder.")
            return "https://picsum.photos/1024"
        except Exception as e:
            print(f"[image_gen_tool] Failed to read local image {path!r}: {e!r}. Falling back to placeholder.")
            return "https://picsum.photos/1024"

    elif mode == "hf":
        # ─────────────────────────────────────────────────────────────────────────
        # NOTE: Hugging Face path intentionally commented out for now to save hits.
        # To re-enable later:
        #   1) uncomment the block below
        #   2) ensure ENV_KEYS['HUGGINGFACE_API_KEY'] is set
        #   3) set IMAGE_GEN_MODE=hf
        # ─────────────────────────────────────────────────────────────────────────
        hf_token = ENV_KEYS.get("HUGGINGFACE_API_KEY")
        if not hf_token:
            print("[image_gen_tool] No HUGGINGFACE_API_KEY; using placeholder.")
            return "https://picsum.photos/1024"

        try:
            # from huggingface_hub import InferenceClient
            # client = InferenceClient(api_key=hf_token)
            # model_to_use = model_id or os.getenv("HF_TTI_MODEL", "black-forest-labs/FLUX.1-schnell")
            #
            # image = client.text_to_image(prompt, model=model_to_use)  # returns PIL.Image
            #
            # buf = BytesIO()
            # image.save(buf, format="PNG")
            # img_bytes = buf.getvalue()
            #
            # # Correct write (your old code used .read on a writable file handle)
            # with open("generated_image.png", "wb") as fp:
            #     fp.write(img_bytes)
            #
            # b64 = base64.b64encode(img_bytes).decode("utf-8")
            # return f"data:image/png;base64,{b64}"
            #
            # For now, explicitly avoid making the call:
            raise RuntimeError("HF generation disabled for testing; set IMAGE_GEN_MODE=hf and uncomment code.")
        except Exception as e:
            print(f"[image_gen_tool] HF generation skipped/failed: {e!r}. Using placeholder.")
            return "https://picsum.photos/1024"

    # Fallback / explicit placeholder mode
    return "https://picsum.photos/1024"




# ===================
# Invocation helpers
# ===================
def _invoke_with_optional_tools(model, messages, tools):
    try:
        if tools:
            return model.invoke(messages, tools=tools)
        return model.invoke(messages)
    except TypeError:
        if tools:
            try:
                return model.invoke(messages, tool_config={"tools": tools})
            except TypeError:
                pass
        return model.invoke(messages)

def _stream_with_optional_tools(model, messages, tools):
    try:
        if tools:
            return model.stream(messages, tools=tools)
        return model.stream(messages)
    except TypeError:
        if tools:
            try:
                return model.stream(messages, tool_config={"tools": tools})
            except TypeError:
                pass
        return model.stream(messages)

_JSON_LINE_RE = re.compile(r'^\s*\{.*\}\s*$', re.DOTALL)

def _maybe_parse_tool_json(first_line: str) -> Optional[Dict[str, Any]]:
    try:
        obj = json.loads(first_line.strip())
        if isinstance(obj, dict) and obj.get("tool") == "generate_image":
            return obj
    except Exception:
        pass
    return None

def _extract_first_line(text: str) -> str:
    # Handle codefence or plain first line
    t = text.strip()
    if t.startswith("```"):
        # try to pull first JSON object inside fence
        # Simple heuristic: grab first non-empty line inside
        lines = [ln for ln in t.splitlines() if ln.strip()]
        if len(lines) >= 2 and lines[1].startswith("{"):
            return lines[1]
    # otherwise just first line
    return t.splitlines()[0] if "\n" in t else t

# ===================
# Public inference APIs
# ===================
def my_genai_chat_function_stream(payload: ChatRequest) -> Generator[str, None, None]:
    """
    Streaming chat generation.
    If image JSON directive might be used, we fall back to single-shot (non-stream) to capture the directive.
    """
    if not payload.message.text and not payload.message.files:
        yield "No message content provided."
        return
    try:
        model = get_model(payload.model_name)
        messages = build_conversation_messages(payload.history, payload, payload.system_prompt)
        if not messages:
            yield "No valid content to process."
            return

        tools, policy, image_json_enabled = get_gemini_tools_and_policy(payload.model_name, payload.message.text or "")
        if policy:
            messages = _insert_policy_message(messages, policy)

        # If we asked for image JSON directives, do a single non-stream call to parse it
        if image_json_enabled:
            resp = _invoke_with_optional_tools(model, messages, tools)
            content = getattr(resp, "content", "") or ""
            if isinstance(content, list):
                content = " ".join(str(p) for p in content if p)

            # Check for first-line JSON directive
            first = _extract_first_line(content)
            tool_req = _maybe_parse_tool_json(first)
            if tool_req:
                img_url = image_gen_tool(tool_req.get("prompt", ""), tool_req.get("model"))
                # Compose final answer: include image and the rest of the model reply (minus the JSON line)
                remaining = "\n".join(content.splitlines()[1:]).strip()
                final = f"![generated]({img_url})\n\n{remaining or ''}".strip()
                yield final or "No response from model."
                return

            # No directive; just stream the original content in one go
            yield content or "No response from model."
            return

        # Normal streaming (no image directive expected)
        for chunk in _stream_with_optional_tools(model, messages, tools):
            if getattr(chunk, "content", None):
                if isinstance(chunk.content, list):
                    yield " ".join(str(p) for p in chunk.content if p)
                else:
                    yield chunk.content
    except Exception as e:
        print(f"Error in LangChain stream call: {e}")
        yield f"Error generating response: {e}"

def my_genai_chat_function(payload: ChatRequest) -> str:
    """Non-streaming variant—Gemini decides search tool; image-gen via JSON directive on first line."""
    if not payload.message.text and not payload.message.files:
        return "No message content provided."
    try:
        model = get_model(payload.model_name)
        messages = build_conversation_messages(payload.history, payload, payload.system_prompt)
        if not messages:
            return "No valid content to process."

        tools, policy, image_json_enabled = get_gemini_tools_and_policy(payload.model_name, payload.message.text or "")
        if policy:
            messages = _insert_policy_message(messages, policy)

        resp = _invoke_with_optional_tools(model, messages, tools)
        content = getattr(resp, "content", "") or ""
        if isinstance(content, list):
            content = " ".join(str(p) for p in content if p)

        if image_json_enabled and content:
            first = _extract_first_line(content)
            tool_req = _maybe_parse_tool_json(first)
            if tool_req:
                img_url = image_gen_tool(tool_req.get("prompt", ""), tool_req.get("model"))
                remaining = "\n".join(content.splitlines()[1:]).strip()
                final = f"![generated]({img_url})\n\n{remaining or ''}".strip()
                return final or "No response from model."

        return content or "No response from model."
    except Exception as e:
        print(f"Error in LangChain call: {e}")
        return f"Error generating response: {e}"
