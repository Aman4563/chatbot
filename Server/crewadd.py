# model.py

from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_mistralai import ChatMistralAI
from langchain_openai import ChatOpenAI
from langchain_anthropic import ChatAnthropic
from langchain_google_genai import ChatGoogleGenerativeAI
import google.generativeai as genai
from google.ai.generativelanguage_v1beta.types import Tool as GenAITool
from schemas import ChatRequest, FileData, Message as APIMessage
import dotenv
import os
from typing import Generator, List, Dict, Any, Union, Optional
import base64
import PyPDF2
import docx
import io
import json
import csv
from pathlib import Path
import os
import base64
import requests
from typing import List
from huggingface_hub import InferenceClient
from io import BytesIO

# NEW: Import for CrewAI integration
from trip_planner import run_trip_planner
from langchain.tools import Tool

dotenv.load_dotenv()

# Model configuration (unchanged)
MODEL_CONFIG = {
    "mistral-large": {
        "provider": "mistral",
        "class": ChatMistralAI,
        "model": "mistral-large-latest",
        "supports_vision": False,
        "supports_files": True,
        "api_key": os.getenv("MISTRAL_API_KEY")
    },
    "mistral-small": {
        "provider": "mistral",
        "class": ChatMistralAI,
        "model": "mistral-small-latest",
        "supports_vision": False,
        "supports_files": True,
        "api_key": os.getenv("MISTRAL_API_KEY")
    },
    "gpt-4o": {
        "provider": "openai",
        "class": ChatOpenAI,
        "model": "gpt-4o",
        "supports_vision": True,
        "supports_files": True,
        "api_key": os.getenv("OPENAI_API_KEY")
    },
    "gpt-4o-mini": {
        "provider": "openai",
        "class": ChatOpenAI,
        "model": "gpt-4o-mini",
        "supports_vision": True,
        "supports_files": True,
        "api_key": os.getenv("OPENAI_API_KEY")
    },
    "claude-3-5-sonnet": {
        "provider": "anthropic",
        "class": ChatAnthropic,
        "model": "claude-3-5-sonnet-20241022",
        "supports_vision": True,
        "supports_files": True,
        "api_key": os.getenv("ANTHROPIC_API_KEY")
    },
    "gemini-1.5-pro": {
        "provider": "google",
        "class": ChatGoogleGenerativeAI,
        "model": "gemini-1.5-pro",
        "supports_vision": True,
        "supports_files": True,
        "api_key": os.getenv("GOOGLE_API_KEY")
    },
    "gemini-1.5-flash": {
        "provider": "google",
        "class": ChatGoogleGenerativeAI,
        "model": "gemini-1.5-flash",
        "supports_vision": True,
        "supports_files": True,
        "api_key": os.getenv("GOOGLE_API_KEY")
    },
    "gemini-2.0-flash": {
        "provider": "google",
        "class": ChatGoogleGenerativeAI,
        "model": "gemini-2.0-flash",
        "supports_vision": True,
        "supports_files": True,
        "api_key": os.getenv("GOOGLE_API_KEY")
    },
    "gemini-2.5-flash": {
        "provider": "google",
        "class": ChatGoogleGenerativeAI,
        "model": "gemini-2.5-flash",
        "supports_vision": True,
        "supports_files": True,
        "api_key": os.getenv("GOOGLE_API_KEY")
    },
    "gemini-2.5-pro": {
        "provider": "google",
        "class": ChatGoogleGenerativeAI,
        "model": "gemini-2.5-pro",
        "supports_vision": True,
        "supports_files": True,
        "api_key": os.getenv("GOOGLE_API_KEY")
    },
}

# --- 2a) Web-search tool via googlesearch-py (scraping) --- (unchanged)
def web_search_tool(query: str, num_results: int = 5) -> List[str]:
    """
    Simple web-search using googlesearch. Returns only absolute URLs.
    """
    try:
        from googlesearch import search
    except ImportError:
        raise RuntimeError("pip install googlesearch-python")
    results: List[str] = []
    for url in search(query):  # no num/stop/pause args
        # skip any relative or non-http links
        if not url.lower().startswith(("http://", "https://")):
            continue
        results.append(url)
        if len(results) >= num_results:
            break
    return results

# --- Add to model.py: Helper for realtime detection and Gemini tool assignment --- (unchanged)
def needs_realtime_data(prompt: str) -> bool:
    """Heuristic check: does the prompt require real-time/search info?"""
    REALTIME_KEYWORDS = [
        "latest", "current", "today", "now", "recent", "news", "score",
        "weather", "who won", "president", "stock price", "exchange rate",
        "trending", "this week", "this month", "2025", "just announced",
        "breaking", "live", "just happened", "just now"
    ]
    prompt_lc = prompt.lower()
    return any(kw in prompt_lc for kw in REALTIME_KEYWORDS)

# NEW: Define CrewAI as a LangChain tool
crewai_trip_tool = Tool(
    name="TripPlanner",
    description="Use this to plan trips based on user queries. Input: full query string. Output: detailed itinerary. Optionally accepts images for analysis (e.g., location photos).",
    func=run_trip_planner  # This calls your CrewAI function
)

# UPDATED: Extend to include CrewAI tool conditionally
def get_gemini_tools_if_needed(model_name: str, prompt: str):
    """Return Gemini Google Search tool if it's a Gemini model and prompt needs real-time data. Also add CrewAI tool for trip planning."""
    tools = []
    # NEW: Add CrewAI tool if query seems travel-related (refine heuristics as needed)
    if any(kw in prompt.lower() for kw in ["trip", "plan travel", "book flight", "hotel", "itinerary", "visit", "vacation"]):
        tools.append(crewai_trip_tool)
    # Existing logic for realtime tools
    if model_name.startswith("gemini") and needs_realtime_data(prompt):
        try:
            from google.ai.generativelanguage_v1beta.types import Tool as GenAITool
        except ImportError:
            return []
        tools.append(GenAITool(google_search={}))
    return tools

def image_gen_tool(prompt: str) -> str:
    """
    1) Try Hugging Face’s free Stable Diffusion Inference API over HTTP.
    2) If that fails for any reason, fall back to a random Picsum image.
    Returns either a data-URL (base64 PNG) or an external URL.
    """
    hf_token = os.getenv("HUGGINGFACE_API_KEY")
    if hf_token:
        try:
            print(f"Using Hugging Face token: {hf_token}")
            client = InferenceClient(
                provider="fal-ai",
                api_key=hf_token,
            )
            # api_url = "https://api-inference.huggingface.co/models/runwayml/stable-diffusion-v1-5"
            # headers = {
            # "Authorization": f"Bearer {hf_token}"
            # }
            payload = {"inputs": prompt}
            image = client.text_to_image(
                prompt,
                model="ZB-Tech/Text-to-Image",
            )
            # print(f"Generated image: {image}")
            # Convert PIL.Image to PNG bytes
            buf = BytesIO()
            image.save(buf, format="PNG")
            buf.seek(0)
            img_bytes = buf.getvalue()
            # Base64-encode and return a data URL
            b64 = base64.b64encode(img_bytes).decode("utf-8")
            print(f"Generated image URL: data:image/png;base64,{b64}")
            open("generated_image.png", "wb").write(img_bytes)  # Save for debugging
            return f"data:image/png;base64,{b64}"
            # resp = requests.post(api_url, headers=headers, json=payload, timeout=60)
            # ct = resp.headers.get("Content-Type", "")
            # If HF returns raw image bytes:
            # if resp.status_code == 200 and ct.startswith("image"):
            # img_bytes = resp.content
            # b64 = base64.b64encode(img_bytes).decode("utf-8")
            # return f"data:image/png;base64,{b64}"
            # HF sometimes returns JSON with an error message:
            # print("HF API returned non-image:", resp.status_code, ct, resp.text)
        except Exception as e:
            print("HuggingFace HTTP API failed:", e)
    # Fallback for UI testing when HF fails or you hit free-tier limits
    return "https://picsum.photos/512"

def get_model(model_name: str):
    """Initialize and return the appropriate LangChain model"""
    if model_name not in MODEL_CONFIG:
        raise ValueError(f"Unsupported model: {model_name}")
    config = MODEL_CONFIG[model_name]
    if not config["api_key"]:
        raise ValueError(f"API key not found for {config['provider']} models")
    # Initialize model based on provider
    if config["provider"] == "mistral":
        return config["class"](
            model=config["model"],
            api_key=config["api_key"],
            temperature=0.7,
            max_tokens=8192,
            streaming=True
        )
    elif config["provider"] == "openai":
        return config["class"](
            model=config["model"],
            api_key=config["api_key"],
            temperature=0.7,
            max_tokens=8192,
            streaming=True
        )
    elif config["provider"] == "anthropic":
        return config["class"](
            model=config["model"],
            api_key=config["api_key"],
            temperature=0.7,
            max_tokens=8192,
            streaming=True
        )
    elif config["provider"] == "google":
        return config["class"](
            model=config["model"],
            google_api_key=config["api_key"],
            temperature=0.7,
            max_output_tokens=8192,
            streaming=True
        )
    else:
        raise ValueError(f"Unknown provider: {config['provider']}")

def process_image_for_langchain(file_data: FileData) -> Dict[str, Any]:
    """Process image for LangChain multimodal input"""
    try:  # For models that support vision, return image data
        return {
            "type": "image_url",
            "image_url": {
                "url": f"data:{file_data.mime_type};base64,{file_data.data}"
            }
        }
    except Exception as e:
        print(f"Error processing image {file_data.filename}: {e}")
        return None

def extract_pdf_content(file_data: FileData) -> str:
    """Extract text content from PDF file"""
    try:
        decoded_data = base64.b64decode(file_data.data)
        pdf_file = io.BytesIO(decoded_data)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text_content = ""
        for page_num, page in enumerate(pdf_reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text_content += f"\n--- Page {page_num} ---\n{page_text}\n"
            except Exception as e:
                text_content += f"\n--- Page {page_num} ---\n[Error extracting page content: {str(e)}]\n"
        if not text_content.strip():
            return "PDF file appears to be empty or contains only images/non-extractable content."
        return text_content.strip()
    except Exception as e:
        return f"Error extracting PDF content: {str(e)}"

def extract_docx_content(file_data: FileData) -> str:
    """Extract text content from DOCX file"""
    try:
        decoded_data = base64.b64decode(file_data.data)
        docx_file = io.BytesIO(decoded_data)
        doc = docx.Document(docx_file)
        text_content = ""
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content += paragraph.text + "\n"
        # Extract tables
        for table in doc.tables:
            text_content += "\n--- Table ---\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                text_content += " | ".join(row_text) + "\n"
        if not text_content.strip():
            return "Word document appears to be empty."
        return text_content.strip()
    except Exception as e:
        return f"Error extracting Word document content: {str(e)}"

def extract_csv_content(file_data: FileData) -> str:
    """Extract and format CSV content"""
    try:
        decoded_data = base64.b64decode(file_data.data)
        csv_content = decoded_data.decode('utf-8')
        # Parse CSV
        csv_file = io.StringIO(csv_content)
        csv_reader = csv.reader(csv_file)
        formatted_content = "CSV Data:\n"
        for row_num, row in enumerate(csv_reader, 1):
            if row_num == 1:
                formatted_content += "Headers: " + " | ".join(row) + "\n"
                formatted_content += "-" * 50 + "\n"
            else:
                formatted_content += f"Row {row_num-1}: " + " | ".join(row) + "\n"
            # Limit to first 100 rows for performance
            if row_num > 100:
                formatted_content += f"\n... (showing first 100 rows, total rows may be more)\n"
                break
        return formatted_content
    except Exception as e:
        return f"Error processing CSV: {str(e)}"

def extract_json_content(file_data: FileData) -> str:
    """Extract and format JSON content"""
    try:
        decoded_data = base64.b64decode(file_data.data)
        json_content = decoded_data.decode('utf-8')
        # Parse and pretty-print JSON
        json_data = json.loads(json_content)
        formatted_json = json.dumps(json_data, indent=2, ensure_ascii=False)
        return f"JSON Data:\n{formatted_json}"
    except json.JSONDecodeError as e:
        return f"Invalid JSON format: {str(e)}"
    except Exception as e:
        return f"Error processing JSON: {str(e)}"

def process_document_for_langchain(file_data: FileData) -> str:
    """Process document for text analysis with proper content extraction"""
    try:
        print(f"Processing document: {file_data.filename} (Type: {file_data.mime_type})")
        if file_data.mime_type == 'text/plain':
            decoded_data = base64.b64decode(file_data.data)
            content = decoded_data.decode('utf-8')
            return f"Text Document: {file_data.filename}\n\nContent:\n{content}"
        elif file_data.mime_type == 'text/csv':
            content = extract_csv_content(file_data)
            return f"CSV Document: {file_data.filename}\n\n{content}"
        elif file_data.mime_type == 'application/json':
            content = extract_json_content(file_data)
            return f"JSON Document: {file_data.filename}\n\n{content}"
        elif file_data.mime_type == 'application/pdf':
            content = extract_pdf_content(file_data)
            return f"PDF Document: {file_data.filename}\n\nExtracted Content:\n{content}"
        elif file_data.mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            content = extract_docx_content(file_data)
            return f"Word Document: {file_data.filename}\n\nExtracted Content:\n{content}"
        elif file_data.mime_type == 'application/msword':
            return f"Legacy Word Document: {file_data.filename}\n\n[Legacy .doc format requires additional processing - please convert to .docx format for full analysis]"
        else:
            return f"Document: {file_data.filename} (Type: {file_data.mime_type})\n[Unsupported document type for content extraction]"
    except Exception as e:
        print(f"Error processing document {file_data.filename}: {e}")
        return f"Error processing document {file_data.filename}: {str(e)}"

def build_conversation_messages(history: List[APIMessage], current_message: ChatRequest, system_prompt: str) -> List:
    """Build conversation messages for LangChain"""
    messages = []  # Add system message
    if system_prompt:
        messages.append(SystemMessage(content=system_prompt))
    # Add conversation history
    for msg in history:
        if msg.role == 'user':
            content = []
            # Add text content
            if msg.content:
                content.append({"type": "text", "text": msg.content})
            # Add files if any
            if hasattr(msg, 'files') and msg.files:
                for file_data in msg.files:
                    if file_data.mime_type.startswith('image/'):
                        # For vision-capable models
                        image_content = process_image_for_langchain(file_data)
                        if image_content:
                            content.append(image_content)
                    else:
                        # For documents, add as text
                        doc_content = process_document_for_langchain(file_data)
                        content.append({"type": "text", "text": doc_content})
            if content:
                if len(content) == 1 and content[0].get("type") == "text":
                    # Simple text message
                    messages.append(HumanMessage(content=content[0]["text"]))
                else:
                    # Multi-modal message
                    messages.append(HumanMessage(content=content))
        elif msg.role == 'assistant':
            messages.append(AIMessage(content=msg.content))
    # Add current message
    current_content = []
    # Add text content
    if current_message.message.text:
        current_content.append({"type": "text", "text": current_message.message.text})
    # Process attached files
    for file_data in current_message.message.files:
        if file_data.mime_type.startswith('image/'):
            image_content = process_image_for_langchain(file_data)
            if image_content:
                current_content.append(image_content)
            print(f"Added image for vision analysis: {file_data.filename}")
        else:
            doc_content = process_document_for_langchain(file_data)
            current_content.append({"type": "text", "text": doc_content})
            print(f"Added document for analysis: {file_data.filename}")
    # Add current message to conversation
    if current_content:
        if len(current_content) == 1 and current_content[0].get("type") == "text":
            # Simple text message
            messages.append(HumanMessage(content=current_content[0]["text"]))
        else:
            # Multi-modal message
            messages.append(HumanMessage(content=current_content))
    return messages

def my_genai_chat_function_stream(payload: ChatRequest) -> Generator[str, None, None]:
    """Enhanced streaming function with LangChain, Gemini tool-calling"""
    if not payload.message.text and not payload.message.files:
        yield "No message content provided."
        return
    try:
        print(f"Processing message: {payload.message.text}")
        print(f"Using model: {payload.model_name}")
        print(f"Files attached: {len(payload.message.files)}")
        print(f"History length: {len(payload.history)}")
        # Get the model
        model = get_model(payload.model_name)
        # Build conversation messages
        messages = build_conversation_messages(
            payload.history,
            payload,
            payload.system_prompt
        )
        if not messages:
            yield "No valid content to process."
            return
        print("Starting stream response...")
        print(f"Messages structure: {len(messages)} messages")
        # UPDATED: Pass tools (including CrewAI if applicable) to the model
        tools = get_gemini_tools_if_needed(payload.model_name, payload.message.text)
        stream_args = {'input': messages}
        if tools:
            stream_args['tools'] = tools
        for chunk in model.stream(**stream_args):
            if hasattr(chunk, 'content') and chunk.content:
                yield chunk.content
        print("Stream response completed")
    except Exception as e:
        print(f"Error in LangChain API call: {e}")
        yield f"Error generating response: {str(e)}"

def my_genai_chat_function(payload: ChatRequest) -> str:
    """Non-streaming version with LangChain, Gemini tool-calling"""
    if not payload.message.text and not payload.message.files:
        return "No message content provided."
    try:
        # Get the model
        model = get_model(payload.model_name)
        # Build conversation messages
        messages = build_conversation_messages(
            payload.history,
            payload,
            payload.system_prompt
        )
        if not messages:
            return "No valid content to process."
        # UPDATED: Pass tools (including CrewAI if applicable) to the model
        tools = get_gemini_tools_if_needed(payload.model_name, payload.message.text)
        invoke_args = {'input': messages}
        if tools:
            invoke_args['tools'] = tools
        response = model.invoke(**invoke_args)
        if response and hasattr(response, 'content') and response.content:
            return response.content
        else:
            return "No response from model."
    except Exception as e:
        print(f"Error in LangChain API call: {e}")
        return f"Error generating response: {str(e)}"
